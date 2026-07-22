# -*- coding: utf-8 -*-
import io
import json
import logging
import requests
from docx import Document

OLLAMA_URL = "http://localhost:11434/api/generate"


def consultar_ollama(prompt):
    payload = {
        "model": "llama3",
        "prompt": prompt,
        "stream": False
    }
    try:
        res = requests.post(OLLAMA_URL, json=payload, timeout=120)
        res.raise_for_status()
        return res.json().get("response", "")
    except Exception as e:
        logging.error(f"Error conectando a Ollama: {e}")
        return "Error: No se pudo generar la redacción cualitativa. Verifique el servicio local de IA."


def generar_documento_word(engine, data):
    """
    Genera documento DOCX integrando datos transaccionales y redacción generada por IA local.
    """
    doc = Document()
    doc.add_heading(f"Informe Mensual de Ejecución - {data.get('nombre_ese', 'ESE')}", 0)

    doc.add_paragraph(f"Resolución: {data.get('resolucion', '')}")
    doc.add_paragraph(f"Mes de Ejecución: {data.get('mes', '')}")
    doc.add_paragraph(f"Representante Legal: {data.get('representante', '')}")

    prompt = (
        f"Actúa como un analista experto en salud pública colombiana. "
        f"Redacta una introducción ejecutiva y profesional de máximo 2 párrafos para un informe mensual "
        f"de ejecución de recursos APS (Resolución 2361/2016) para la ESE '{data.get('nombre_ese')}' "
        f"en los municipios de '{data.get('municipios')}'. Enfatiza el cumplimiento de metas."
    )

    introduccion_ia = consultar_ollama(prompt)

    doc.add_heading("1. Introducción Cualitativa (Generada por IA)", level=1)
    doc.add_paragraph(introduccion_ia)

    doc.add_heading("2. Consolidado Financiero Transaccional", level=1)
    doc.add_paragraph(
        "A continuación se presenta el resumen de control financiero extraído directamente de la base de datos de auditoría SER124:")

    # Tabla Dummy simulando extracción de BD para el informe
    tabla = doc.add_table(rows=1, cols=3)
    tabla.style = 'Table Grid'
    hdr_cells = tabla.rows[0].cells
    hdr_cells[0].text = 'Concepto'
    hdr_cells[1].text = 'Cantidad'
    hdr_cells[2].text = 'Valor'

    row_cells = tabla.add_row().cells
    row_cells[0].text = 'Actos de Incorporación'
    row_cells[1].text = '1'
    row_cells[2].text = '$ Consultar BD'

    nombre_archivo = f"Informe_Mensual_{data.get('nit_ese', 'NIT')}_{data.get('mes', 'Mes').replace(' ', '_')}.docx"

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    return nombre_archivo, stream